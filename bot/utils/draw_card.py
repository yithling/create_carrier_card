from datetime import datetime
from dateutil.relativedelta import relativedelta

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm

from bot.config import RESULT_FILE_PATH, QR_FILE_PATH


def check_osgop_data(osgop_data) -> bool:
    if osgop_data is None:
        return False
    return True


def build_mo_card(
        organization_name, 
        driver_name, 
        phone_number, 
        car_number, 
        legal_address, 
        license_data, 
        osgop_data,
        file_path
) -> str:

    #Создаем карточку для Москвы
    #Загружаем таблицу
    doc = Document(file_path)
    tables = doc.tables
    #Указываем размер для QR изображения
    width = Cm(2.83)
    height = Cm(2.83)
    if tables:
        table = tables[0]
        #Проходим по каждой записи
        for row in table.rows:
            for cell in row.cells:
                #Если в таблице есть ключевые слова, то мы их заменяем на те даныне, которые получили
                if cell.text == "organization":
                    organization = organization_name.split()
                    if organization[0] == "ИП":
                        cell.text = "Индивидуальный предприниматель"
                    elif organization[0] == "ООО":
                        cell.text = "ОБЩЕСТВО С ОГРАНИЧЕННОЙ ОТВЕТСТВЕННОСТЬЮ"
                    else:
                        cell.text = "Самозанятый"

                if cell.text == "name_organization":
                    cell.text = organization_name

                if cell.text == "driver_name":
                    cell.text = driver_name

                if cell.text == "legal_address":
                    cell.text = legal_address

                if cell.text == "phone_number":
                    cell.text = phone_number

                if cell.text == "car_number":
                    cell.text = car_number

                if cell.text == "osgop_company" and check_osgop_data(osgop_data):
                    cell.text = osgop_data["Наименование страховщика"]

                if cell.text == "osgop_number" and check_osgop_data(osgop_data):
                    cell.text = osgop_data["№"].lstrip()

                if cell.text == "validity_period_from":
                    cell.text = license_data["Дата предоставления разрешения:"]

                if cell.text == "valid_until":
                    cell.text = license_data["Срок действия разрешения:"]

                if cell.text == "osgop_validity_period_from" and check_osgop_data(osgop_data):
                    cell.text = osgop_data["Дата начала ответственности"]

                if cell.text == "osgop_valid_until" and check_osgop_data(osgop_data):
                    date_from = str(osgop_data["Дата начала ответственности"]).strip()
                    date_obj = datetime.strptime(date_from, "%d.%m.%Y")
                    new_date = date_obj + relativedelta(years=1)
                    result_date = new_date - relativedelta(days=1)
                    date_until = result_date.strftime("%d.%m.%Y")
                    cell.text  = date_until

                if cell.text == "QR_image":
                    cell.text = ""
                    paragraph = cell.paragraphs[0]
                    run = paragraph.add_run()
                    run.add_picture(QR_FILE_PATH, width=width, height=height)

                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if "Такси car_license_number" in paragraph.text:
                        paragraph.text = paragraph.text.replace("Такси car_license_number", 
                                                                f"Такси {license_data["Номер реестровой записи в региональном реестре легкового такси:"]}"
                                                                )

                    if "Перевозчик  carier_license_number" in paragraph.text:
                        paragraph.text = paragraph.text.replace("Перевозчик  carier_license_number",
                                                                    f"Перевозчик {license_data["Номер реестровой записи в региональном реестре перевозчиков легковым такси:"]}")

                    if "osgop_company" in paragraph.text and check_osgop_data(osgop_data):
                        paragraph.text = paragraph.text.replace("osgop_company",
                                                                osgop_data["Наименование страховщика"])

                    if "osgop_number" in paragraph.text and check_osgop_data(osgop_data):
                        paragraph.text = paragraph.text.replace("№  osgop_number",
                                                                f"№ {osgop_data["№"]}")

                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    #Сохраняем файл и возвращаем путь до файла
    file_path = RESULT_FILE_PATH.format(driver_name, car_number)
    doc.save(file_path)
    return file_path


def build_msk_card(
        organization_name, 
        driver_name, 
        phone_number, 
        car_number, 
        legal_address, 
        license_data, 
        osgop_data,
        file_path
) -> str:

    #Создаем карточку для Москвы
    #Загружаем таблицу
    doc = Document(file_path)
    tables = doc.tables
    if tables:
        table = tables[0]
        #Проходимся по каждой записи
        for row in table.rows:
            for cell in row.cells:
                #Если в таблице есть ключевые слова, то мы их заменяем на те даныне, которые получили
                if cell.text == "organization":
                    organization = organization_name.split()
                    if organization[0] == "ИП":
                        cell.text = "Индивидуальный предприниматель"
                    elif organization[0] == "ООО":
                        cell.text = "ОБЩЕСТВО С ОГРАНИЧЕННОЙ ОТВЕТСТВЕННОСТЬЮ"
                    else:
                        cell.text = "Самозанятый"

                if cell.text == "name_organization":
                    cell.text = organization_name

                if cell.text == "driver_name":
                    cell.text = driver_name

                if cell.text == "legal_address":
                    cell.text = legal_address

                if cell.text == "phone_number":
                    cell.text = phone_number

                if cell.text == "car_number":
                    cell.text = car_number

                if cell.text == "osgop_company" and check_osgop_data(osgop_data):
                    cell.text = osgop_data["Наименование страховщика"]

                if cell.text == "osgop_number" and check_osgop_data(osgop_data):
                    cell.text = osgop_data["№"].lstrip()

                if cell.text == "validity_period_from":
                    cell.text = license_data["Дата внесения записи в региональный реестр перевозчиков легковым такси"]

                if cell.text == "valid_until":
                    cell.text = license_data["Дата окончания срока действия разрешения"]

                if cell.text == "osgop_validity_period_from" and check_osgop_data(osgop_data):
                    cell.text = osgop_data["Дата начала ответственности"]

                if cell.text == "osgop_valid_until" and check_osgop_data(osgop_data):
                    date_from = str(osgop_data["Дата начала ответственности"]).strip()
                    date_obj = datetime.strptime(date_from, "%d.%m.%Y")
                    new_date = date_obj + relativedelta(years=1)
                    result_date = new_date - relativedelta(days=1)
                    date_until = result_date.strftime("%d.%m.%Y")
                    cell.text  = date_until

                for paragraph in cell.paragraphs:
                    #Делаем по центру
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    if "Такси car_license_number" in paragraph.text:
                        new_value = f"Такси {license_data["Номер записи в региональном реестре легковых такси, содержащий сведения о легковом такси"]}"
                        paragraph.text = paragraph.text.replace("Такси car_license_number", new_value)

                    if "Перевозчик  carier_license_number" in paragraph.text:
                        paragraph.text = paragraph.text.replace("Перевозчик  carier_license_number",
                                                                    f"Перевозчик {license_data["Номер записи в региональном реестре перевозчиков легковым такси, содержащий сведения о перевозчике"]}")

                    if "osgop_company" in paragraph.text and check_osgop_data(osgop_data):
                        paragraph.text = paragraph.text.replace("osgop_company",
                                                                osgop_data["Наименование страховщика"])

                    if "osgop_number" in paragraph.text and check_osgop_data(osgop_data):
                        paragraph.text = paragraph.text.replace("№  osgop_number",
                                                                f"№ {osgop_data["№"]}")

                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    #Сохраняем файл и возвращаем путь до файла
    file_path = RESULT_FILE_PATH.format(driver_name, car_number)
    doc.save(file_path)
    return file_path
